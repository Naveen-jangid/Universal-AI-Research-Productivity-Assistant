"""
Document ingestion pipeline.
Handles: PDF / TXT / DOCX / MD extraction → chunking → embedding → FAISS storage.
"""

import logging
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from backend.core.config import settings
from backend.core.database import register_document, update_document_status
from backend.vectorstore.faiss_store import add_documents

logger = logging.getLogger(__name__)


# ── Text extraction ─────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF

        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text("text"))
        return "\n".join(text_parts)
    except ImportError:
        # Fallback to pypdf
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def extract_text(file_path: str) -> str:
    """
    Route to the correct extractor based on file extension.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        Raw extracted text.

    Raises:
        ValueError: If the file type is not supported.
    """
    suffix = Path(file_path).suffix.lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".doc": extract_text_from_docx,
        ".txt": lambda p: Path(p).read_text(encoding="utf-8", errors="replace"),
        ".md": lambda p: Path(p).read_text(encoding="utf-8", errors="replace"),
        ".rst": lambda p: Path(p).read_text(encoding="utf-8", errors="replace"),
    }
    if suffix not in extractors:
        raise ValueError(
            f"Unsupported document type '{suffix}'. "
            f"Supported: {list(extractors.keys())}"
        )
    text = extractors[suffix](file_path)
    logger.info("Extracted %d chars from '%s'", len(text), Path(file_path).name)
    return text


# ── Chunking ────────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    metadata: Optional[Dict] = None,
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[Document]:
    """
    Split text into overlapping chunks, returning LangChain Document objects.

    Args:
        text: Full extracted text.
        metadata: Arbitrary metadata to attach to every chunk.
        chunk_size: Characters per chunk (defaults to settings.CHUNK_SIZE).
        chunk_overlap: Overlap between chunks (defaults to settings.CHUNK_OVERLAP).

    Returns:
        List of Document objects ready for embedding.
    """
    chunk_size = chunk_size or settings.CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""],
    )
    chunks = splitter.split_text(text)

    metadata = metadata or {}
    docs = [
        Document(page_content=chunk, metadata={**metadata, "chunk_index": i})
        for i, chunk in enumerate(chunks)
    ]
    logger.info("Created %d chunks (size=%d, overlap=%d)", len(docs), chunk_size, chunk_overlap)
    return docs


# ── Full ingestion pipeline ─────────────────────────────────────────────────

def ingest_document(
    file_path: str,
    doc_id: Optional[str] = None,
    namespace: str = "default",
) -> Tuple[str, int]:
    """
    End-to-end pipeline: extract → chunk → embed → store.

    Args:
        file_path: Path to the uploaded file.
        doc_id: Optional unique identifier. Generated if not provided.
        namespace: FAISS namespace for the vector store.

    Returns:
        Tuple of (doc_id, chunk_count).
    """
    path = Path(file_path)
    doc_id = doc_id or str(uuid.uuid4())

    # Register in SQLite
    register_document(
        doc_id=doc_id,
        filename=path.name,
        file_type=path.suffix.lower(),
        file_size=path.stat().st_size,
        metadata={"namespace": namespace, "file_path": str(path)},
    )

    try:
        # 1. Extract text
        text = extract_text(file_path)
        if not text.strip():
            raise ValueError("Document appears to be empty or unreadable.")

        # 2. Chunk
        metadata = {
            "source": path.name,
            "doc_id": doc_id,
            "namespace": namespace,
        }
        chunks = chunk_text(text, metadata=metadata)

        # 3. Embed + store in FAISS
        added = add_documents(chunks, namespace=namespace)

        # 4. Update DB record
        update_document_status(doc_id, "ready", chunk_count=added)
        logger.info("Document '%s' ingested: %d chunks stored.", path.name, added)
        return doc_id, added

    except Exception as exc:
        update_document_status(doc_id, "error")
        logger.error("Ingestion failed for '%s': %s", path.name, exc, exc_info=True)
        raise
